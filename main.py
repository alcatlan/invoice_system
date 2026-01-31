from mailer import enviar_factura_por_email
from tracker import registrar_en_excel
from docx import Document
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx2pdf import convert
import os 

def obtener_y_actualizar_consecutivo(nombre_archivo="ultimo_numero.txt"):
    if not os.path.exists(nombre_archivo):
        with open(nombre_archivo, "w") as f:
            f.write("1")
        return 1
    
    with open(nombre_archivo, "r") as f:
        numero = int(f.read())
    
    with open(nombre_archivo, "w") as f:
        f.write(str(numero + 1))
        
    return numero

# --- ESTA ES LA FUNCIÓN QUE LLAMARÁ TU APP.PY ---
def generar_factura_completa(TARIFA_HORA, email_cliente, servicios, numero_factura):
    # 1. Capturamos la fecha actual
    fecha_hoy = datetime.now().strftime("%B %d, %Y")
    doc = Document()
    from docx.shared import Inches
    # 1. Accedemos al encabezado de la primera sección
    header = doc.sections[0].header

    # 2. Añadimos un párrafo al encabezado
    p = header.paragraphs[0]

    # 3. Usamos un "run" para insertar la imagen
    run = p.add_run()
    run.add_picture("logo.png", width=Inches(1.5)) # Ajusta el tamaño aquí 📏
    # 4. Alineamos el párrafo a la derecha
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT    

    # 1. Agregamos párrafos vacíos después del encabezado para dar aire 🌬️
    doc.add_paragraph("\n\n")

    # 2. Encabezados de la factura
    doc.add_heading('INVOICE NB - TX3 USER UPDATE', 0)
    doc.add_paragraph("Supplier:")
    doc.add_paragraph("Alejandro Manrique Florez")
    doc.add_paragraph("119  50 Via Calabria St")
    doc.add_paragraph(f"Email: {email_cliente}")

    numero_factura = obtener_y_actualizar_consecutivo()
    invoice_string = f"2026-{numero_factura:03}"
    doc.add_paragraph(f"Invoice No: {invoice_string}")
    doc.add_paragraph(f"Date Issued: {fecha_hoy}")

    # 3. Configuracion de la tabla
    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = 'Table Grid'
    anchos = [Inches(0.7), Inches(3.4), Inches(1.2), Inches(1.2)]

    for i, columna in enumerate(tabla.columns):
        columna.width = anchos[i]

    encabezados = tabla.rows[0].cells
    titulos = ['Hour', 'Description', 'Rate (CAD)', 'Amount (CAD)']

    for i, titulo in enumerate(titulos):
        p = encabezados[i].paragraphs[0]
        p.text = titulo
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        p.runs[0].bold = True 

    # 4. Llenado de la tabla y cálculos (Usando los datos de la GUI)
    subtotal = 0
    # Aseguramos que TARIFA_HORA sea un número
    tarifa_num = float(TARIFA_HORA)

    for s in servicios:
        monto_servicio = s["horas"] * tarifa_num
        subtotal += monto_servicio
        
        nueva_fila = tabla.add_row().cells
        nueva_fila[0].text = f"{s['horas']}h"
        nueva_fila[1].text = s["descripcion"]
        
        p_rate = nueva_fila[2].paragraphs[0]
        p_rate.text = f"${tarifa_num:.2f}"
        p_rate.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        p_monto = nueva_fila[3].paragraphs[0]
        p_monto.text = f"${monto_servicio:.2f}"
        p_monto.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 5. Totales finales
    doc.add_paragraph(f"\nSubtotal: ${subtotal:.2f} CAD").alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("HST: [Not Registered]").alignment = WD_ALIGN_PARAGRAPH.LEFT
    total_final = doc.add_paragraph(f"Total: ${subtotal:.2f} CAD")
    total_final.alignment = WD_ALIGN_PARAGRAPH.LEFT
    total_final.runs[0].bold = True 

    # 6. Guardado y Conversión
    fecha_limpia = fecha_hoy.replace(' ', '_').replace(',', '')
    nombre_unico = f"Invoice_{invoice_string}_{fecha_limpia}"

    doc.save(f"{nombre_unico}.docx")
    print(f"✨ ¡Word generado: {nombre_unico}.docx!")

    print("Generando PDF... ⏳")
    convert(f"{nombre_unico}.docx", f"{nombre_unico}.pdf")
    print(f"✅ ¡PDF creado: {nombre_unico}.pdf!")

    # 7. Excel y Correo Automático
    registrar_en_excel(fecha_hoy, invoice_string, subtotal)
    
    print(f"Enviando factura {invoice_string} a {email_cliente}... ⏳")
    exito = enviar_factura_por_email(email_cliente, f'{nombre_unico}.pdf', invoice_string)
    
    if exito:
        print("✅ ¡Correo enviado exitosamente!")
    else:
        print("❌ Error al enviar el correo.")
    
    return nombre_unico # Devolvemos el nombre por si la GUI lo necesita